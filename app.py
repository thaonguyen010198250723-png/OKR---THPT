import streamlit as st
import pandas as pd
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import json
import unicodedata
import time
import io
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==============================================================================
# 1. CẤU HÌNH & GIAO DIỆN
# ==============================================================================
st.set_page_config(
    page_title="Hệ thống Quản lý OKR (V10 - Admin Fix)",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .stApp { background-color: #fcfcfc; }
    .stButton>button { background-color: #FF8C00; color: white; border-radius: 5px; border: none; }
    .stButton>button:hover { background-color: #e07b00; color: white; }
    h1, h2, h3 { color: #E65100; }
    .status-badge {
        display: inline-block; padding: 5px 10px; border-radius: 15px;
        font-size: 0.8rem; font-weight: bold; text-align: center; min-width: 80px;
    }
    .badge-green { background-color: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
    .badge-red { background-color: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
    .badge-yellow { background-color: #fff3cd; color: #856404; border: 1px solid #ffeeba; }
    .badge-grey { background-color: #e2e3e5; color: #383d41; border: 1px solid #d6d8db; }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 2. XỬ LÝ KẾT NỐI & CORE DATA FUNCTIONS
# ==============================================================================
SHEET_ID = "14E2JfVyOhGMa7T1VA44F31IaPMWIVIPRApo4B-ipDLk"

@st.cache_resource
def init_connection():
    try:
        scope = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
        creds_dict = json.loads(st.secrets["service_account"]["info"])
        creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        client = gspread.authorize(creds)
        return client
    except Exception as e:
        st.error(f"Lỗi kết nối API Google: {e}")
        return None

def get_worksheet(sheet_name):
    client = init_connection()
    if not client: return None
    try:
        sh = client.open_by_key(SHEET_ID)
        return sh.worksheet(sheet_name)
    except gspread.WorksheetNotFound:
        # Tự động tạo tab nếu thiếu
        sh = client.open_by_key(SHEET_ID)
        ws = sh.add_worksheet(title=sheet_name, rows=100, cols=20)
        headers = {
            "Users": ["Email", "Password", "HoTen", "VaiTro", "TenLop"],
            "Classes": ["TenLop", "EmailGVCN", "SiSo"],
            "Periods": ["ID", "TenDot", "TrangThai"],
            "Relationships": ["Email_HocSinh", "Email_PhuHuynh"],
            "OKRs": ["ID", "Email_HocSinh", "ID_Dot", "MucTieu", "KetQuaThenChot", "TienDo", "TrangThai", "NhanXet_GV", "NhanXet_PH", "MinhChung", "TargetValue", "ActualValue", "Unit", "DeleteRequest"],
            "FinalReviews": ["Email_HocSinh", "ID_Dot", "NhanXet_GV", "NhanXet_PH", "DaGui_PH"]
        }
        if sheet_name in headers:
            ws.append_row(headers[sheet_name])
            if sheet_name == "Users":
                ws.append_row(["admin@school.com", "123", "Quản Trị Viên", "Admin", ""])
        return ws
    except Exception as e:
        st.error(f"Lỗi truy cập dữ liệu: {e}")
        return None

@st.cache_data(ttl=10)
def load_data(sheet_name):
    """Đọc dữ liệu an toàn với caching"""
    ws = get_worksheet(sheet_name)
    if not ws: return pd.DataFrame()
    try:
        data = ws.get_all_records()
        df = pd.DataFrame(data)
        
        # Xử lý numeric columns
        numeric_cols = ['ID', 'ID_Dot', 'TargetValue', 'ActualValue', 'DeleteRequest', 'SiSo', 'DaGui_PH']
        for col in numeric_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
        
        # Fix cột Users
        if sheet_name == "Users":
            if 'ClassID' in df.columns and 'TenLop' not in df.columns:
                df = df.rename(columns={'ClassID': 'TenLop'})
            if 'TenLop' not in df.columns: df['TenLop'] = ""
            
            # --- FIX QUAN TRỌNG: Ép kiểu mật khẩu thành chuỗi ---
            if 'Password' in df.columns:
                df['Password'] = df['Password'].astype(str)
            
        return df
    except Exception as e:
        return pd.DataFrame()

def batch_add_records(sheet_name, rows_data):
    """Thêm nhiều dòng cùng lúc (Fix lỗi 429)"""
    ws = get_worksheet(sheet_name)
    if ws and rows_data:
        try:
            ws.append_rows(rows_data, value_input_option='USER_ENTERED')
            st.cache_data.clear()
            return True
        except Exception as e:
            st.error(f"Lỗi Batch Import: {e}")
            return False
    return False

def update_cell_value(sheet_name, match_col, match_val, update_col, update_val, match_col_2=None, match_val_2=None):
    """Cập nhật 1 ô"""
    ws = get_worksheet(sheet_name)
    if not ws: return
    try:
        headers = ws.row_values(1)
        if update_col not in headers: return
        update_col_idx = headers.index(update_col) + 1
        
        # Mapping cột User cũ
        real_match_col = match_col
        if sheet_name == "Users" and match_col == "TenLop" and "ClassID" in headers:
            real_match_col = "ClassID"

        data = ws.get_all_records()
        row_idx = -1
        
        for i, row in enumerate(data):
            val1 = str(row.get(real_match_col, ''))
            is_match = False
            
            if match_col_2:
                val2 = str(row.get(match_col_2, ''))
                if val1 == str(match_val) and val2 == str(match_val_2):
                    is_match = True
            else:
                if val1 == str(match_val):
                    is_match = True
            
            if is_match:
                row_idx = i + 2 
                break
        
        if row_idx != -1:
            ws.update_cell(row_idx, update_col_idx, update_val)
            st.cache_data.clear()
            return True
    except Exception as e:
        st.error(f"Lỗi cập nhật: {e}")
        return False

def delete_record(sheet_name, match_col, match_val):
    ws = get_worksheet(sheet_name)
    if not ws: return
    try:
        # Tìm và xóa dòng đầu tiên khớp (Simple delete)
        cell = ws.find(str(match_val), in_column=ws.find(match_col).col)
        if cell:
            ws.delete_rows(cell.row)
            st.cache_data.clear()
    except: pass

def get_next_id(sheet_name):
    df = load_data(sheet_name)
    if df.empty or 'ID' not in df.columns: return 1
    return int(df['ID'].max()) + 1

# --- CÁC HÀM XỬ LÝ LOGIC PHỨC TẠP ---

def upsert_final_review(email, id_dot, col_name, value):
    """Insert hoặc Update nhận xét"""
    df = load_data("FinalReviews")
    exists = False
    if not df.empty:
        mask = (df['Email_HocSinh'] == email) & (df['ID_Dot'] == id_dot)
        if not df[mask].empty: exists = True
    
    if exists:
        update_cell_value("FinalReviews", "Email_HocSinh", email, col_name, value, "ID_Dot", id_dot)
    else:
        row = [email, id_dot, "", "", 0]
        if col_name == "NhanXet_GV": row[2] = value
        elif col_name == "NhanXet_PH": row[3] = value
        elif col_name == "DaGui_PH": row[4] = value
        batch_add_records("FinalReviews", [row])

def update_student_email_cascade(old_email, new_email):
    """Đổi Email học sinh và cập nhật tất cả bảng liên quan"""
    try:
        # 1. Update Users
        update_cell_value("Users", "Email", old_email, "Email", new_email)
        
        # 2. Update Related Tables
        tables_map = {
            "Relationships": ["Email_HocSinh", "Email_PhuHuynh"],
            "OKRs": ["Email_HocSinh"],
            "FinalReviews": ["Email_HocSinh"]
        }
        
        for table, cols in tables_map.items():
            ws = get_worksheet(table)
            if not ws: continue
            try:
                headers = ws.row_values(1)
                for col_name in cols:
                    if col_name in headers:
                        col_idx = headers.index(col_name) + 1
                        cells = ws.findall(old_email, in_column=col_idx)
                        if cells:
                            for cell in cells: cell.value = new_email
                            ws.update_cells(cells)
            except Exception as ex:
                print(f"Skip {table}: {ex}")
        
        st.cache_data.clear()
        return True
    except Exception as e:
        st.error(f"Lỗi đồng bộ Email: {e}")
        return False

def delete_student_fully(email):
    """Xóa hoàn toàn học sinh và dữ liệu liên quan"""
    try:
        # 1. Delete from Users
        delete_record("Users", "Email", email)
        
        # 2. Delete related data (Tìm và xóa)
        delete_record("Relationships", "Email_HocSinh", email)
        delete_record("FinalReviews", "Email_HocSinh", email)
        
        # OKRs có thể có nhiều dòng, cần xóa hết
        ws_okr = get_worksheet("OKRs")
        if ws_okr:
            try:
                # Tìm tất cả cells chứa email
                cells = ws_okr.findall(email, in_column=ws_okr.find("Email_HocSinh").col)
                # Xóa từ dưới lên để không lệch index
                rows_to_del = sorted([c.row for c in cells], reverse=True)
                for r in rows_to_del:
                    ws_okr.delete_rows(r)
            except: pass
            
        st.cache_data.clear()
        return True
    except Exception as e:
        st.error(f"Lỗi xóa dữ liệu: {e}")
        return False

# ==============================================================================
# 3. TIỆN ÍCH WORD & TÍNH TOÁN
# ==============================================================================

def calculate_percent(actual, target):
    try:
        return round((float(actual) / float(target) * 100), 1) if float(target) != 0 else 0
    except: return 0

def get_rank(percent):
    if percent >= 80: return "Tốt", "green"
    elif percent >= 65: return "Khá", "blue"
    elif percent >= 50: return "Đạt", "orange"
    return "Chưa đạt", "red"

def add_student_report_to_doc(doc, student_name, class_name, period_name, okr_df, review_gv, review_ph):
    doc.add_heading('PHIẾU KẾT QUẢ OKR', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Họ tên: {student_name} | Lớp: {class_name}')
    doc.add_paragraph(f'Đợt: {period_name} | Ngày: {time.strftime("%d/%m/%Y")}')
    
    doc.add_heading('1. Chi tiết Mục tiêu', level=1)
    if not okr_df.empty:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, t in enumerate(['Mục Tiêu', 'Kết Quả (KR)', 'Target', 'Actual', '%']):
            hdr[i].text = t
        
        total = 0
        for _, row in okr_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['MucTieu'])
            cells[1].text = str(row['KetQuaThenChot'])
            cells[2].text = f"{row['TargetValue']} {row['Unit']}"
            cells[3].text = f"{row['ActualValue']} {row['Unit']}"
            pct = calculate_percent(row['ActualValue'], row['TargetValue'])
            cells[4].text = f"{pct}%"
            total += pct
        
        avg = round(total / len(okr_df), 1)
        rank, _ = get_rank(avg)
        doc.add_paragraph(f'\nTrung bình: {avg}% - Xếp loại: {rank}')
    else:
        doc.add_paragraph('Chưa có dữ liệu.')

    doc.add_heading('2. Nhận xét', level=1)
    doc.add_paragraph(f"GVCN: {review_gv if review_gv else '---'}")
    doc.add_paragraph(f"Phụ huynh: {review_ph if review_ph else '---'}")

def create_single_docx(student_name, class_name, period_name, okr_df, review_gv, review_ph):
    doc = Document()
    add_student_report_to_doc(doc, student_name, class_name, period_name, okr_df, review_gv, review_ph)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_class_report_docx(class_name, list_students, all_okrs, all_reviews, period_name, period_id):
    doc = Document()
    count = 0
    for idx, hs in list_students.iterrows():
        count += 1
        hs_okrs = pd.DataFrame()
        if not all_okrs.empty:
            hs_okrs = all_okrs[(all_okrs['Email_HocSinh'] == hs['Email']) & (all_okrs['ID_Dot'] == period_id)]
        
        rev_gv, rev_ph = "", ""
        if not all_reviews.empty:
            r = all_reviews[(all_reviews['Email_HocSinh'] == hs['Email']) & (all_reviews['ID_Dot'] == period_id)]
            if not r.empty:
                rev_gv = r.iloc[0]['NhanXet_GV']
                rev_ph = r.iloc[0]['NhanXet_PH']
        
        add_student_report_to_doc(doc, hs['HoTen'], class_name, period_name, hs_okrs, rev_gv, rev_ph)
        if count < len(list_students): doc.add_page_break()
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def change_password_ui(email):
    with st.expander("🔐 Đổi mật khẩu"):
        with st.form("change_pass"):
            p1 = st.text_input("Mật khẩu mới", type="password")
            p2 = st.text_input("Xác nhận", type="password")
            if st.form_submit_button("Lưu"):
                if p1 == p2 and p1:
                    update_cell_value("Users", "Email", email, "Password", p1)
                    st.success("Thành công!")
                else: st.error("Mật khẩu không khớp.")

def get_periods_map(role):
    df = load_data("Periods")
    if df.empty or 'TrangThai' not in df.columns: return {}
    if role != 'Admin': df = df[df['TrangThai'] == 'Mo']
    if 'TenDot' in df.columns and 'ID' in df.columns:
        return dict(zip(df['TenDot'], df['ID']))
    return {}

# ==============================================================================
# 4. DASHBOARD LOGIC
# ==============================================================================

def login_page():
    st.markdown("<h2 style='text-align: center;'>🔐 Đăng Nhập</h2>", unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        with st.form("login"):
            email = st.text_input("Email")
            pwd = st.text_input("Mật khẩu", type="password")
            if st.form_submit_button("Đăng nhập"):
                
                # --- MASTER KEY (VÀO THẲNG KHÔNG CẦN EXCEL) ---
                if email.strip() == "admin@school.com" and str(pwd).strip() == "123":
                    st.success("Đang đăng nhập quyền Admin...")
                    st.session_state['user'] = {
                        'email': 'admin@school.com', 
                        'name': 'Quản Trị Viên (Gốc)', 
                        'role': 'Admin', 
                        'ten_lop': ''
                    }
                    st.rerun()
                # ----------------------------------------------

                df = load_data("Users")
                if df.empty:
                    st.error("Lỗi kết nối CSDL hoặc File đang bận.")
                    return
                
                # Fix lỗi định dạng số/chữ
                df['Password'] = df['Password'].astype(str)
                user = df[(df['Email'] == email) & (df['Password'] == str(pwd))]
                
                if not user.empty:
                    u = user.iloc[0]
                    st.session_state['user'] = {
                        'email': u['Email'], 'name': u['HoTen'], 
                        'role': u['VaiTro'], 'ten_lop': u.get('TenLop', '')
                    }
                    st.rerun()
                else: st.error("Sai thông tin.")
    
    with col2:
        st.info("""
        **Thông tin đăng nhập mặc định:**
        
        🔑 **Admin:** admin@school.com | Pass: 123
        
        *(Tài khoản Admin này luôn hoạt động kể cả khi mất kết nối)*
        """)

# --- ADMIN ---
def admin_dashboard(period_id):
    st.header("🛠️ Admin Dashboard")
    change_password_ui(st.session_state['user']['email'])
    
    tab1, tab2, tab3 = st.tabs(["Thống kê Lớp", "Quản lý User", "Quản lý Đợt"])
    
    with tab1:
        st.subheader(f"📊 Thống kê - Đợt ID: {period_id}")
        classes = load_data("Classes")
        
        if not classes.empty:
            reviews = load_data("FinalReviews")
            okrs = load_data("OKRs")
            users = load_data("Users")
            
            stats = []
            for _, cl in classes.iterrows():
                ten_lop = cl.get('TenLop', '')
                siso = int(cl.get('SiSo', 0))
                
                hs_list = []
                if not users.empty and 'TenLop' in users.columns:
                    hs_list = users[users['TenLop'] == ten_lop]['Email'].tolist()
                
                okr_count = 0
                approved_count = 0
                if hs_list:
                    if not okrs.empty:
                        okr_count = okrs[(okrs['ID_Dot'] == period_id) & (okrs['Email_HocSinh'].isin(hs_list))].shape[0]
                    if not reviews.empty:
                        approved_count = reviews[(reviews['ID_Dot'] == period_id) & (reviews['Email_HocSinh'].isin(hs_list)) & (reviews['NhanXet_GV'] != "")].shape[0]
                
                stats.append({
                    "Lớp": ten_lop, "GVCN": cl.get('EmailGVCN', ''), "Sĩ số": siso,
                    "Tổng OKR": okr_count, "HS Đã Duyệt": f"{approved_count}/{len(hs_list)}"
                })
            
            df_stats = pd.DataFrame(stats)
            st.dataframe(df_stats)
            if not df_stats.empty: st.bar_chart(df_stats.set_index("Lớp")[["Tổng OKR"]])

        st.divider()
        with st.form("create_class"):
            c1, c2, c3 = st.columns(3)
            name = c1.text_input("Tên Lớp")
            gv = c2.text_input("Email GVCN")
            ss = c3.number_input("Sĩ số", 30)
            if st.form_submit_button("Tạo Lớp"):
                try:
                    batch_add_records("Classes", [[name, gv, ss]])
                    all_u = load_data("Users")
                    if all_u.empty or gv not in all_u['Email'].values:
                        batch_add_records("Users", [[gv, "123", f"GV {name}", "GiaoVien", ""]])
                    st.success("Xong!")
                    st.rerun()
                except Exception as e: st.error(str(e))

    with tab2:
        search = st.text_input("Tìm Email:")
        if search:
            u = load_data("Users")
            if not u.empty:
                res = u[u['Email'] == search]
                st.write(res)
                if not res.empty and st.button("Reset Pass 123"):
                    update_cell_value("Users", "Email", search, "Password", "123")
                    st.success("Đã reset.")

    with tab3:
        periods = load_data("Periods")
        for _, row in periods.iterrows():
            c1, c2 = st.columns([4, 1])
            is_open = row.get('TrangThai') == 'Mo'
            p_id = row.get('ID')
            p_name = row.get('TenDot', '')
            toggle = c1.toggle(f"{p_name}", value=is_open, key=f"p_{p_id}")
            if toggle != is_open:
                update_cell_value("Periods", "ID", p_id, "TrangThai", "Mo" if toggle else "Khoa")
                st.rerun()
            if c2.button("🗑️", key=f"del_p_{p_id}"):
                delete_record("Periods", "ID", p_id)
                st.rerun()
        
        with st.form("add_p"):
            pn = st.text_input("Tên đợt")
            if st.form_submit_button("Thêm"):
                nid = get_next_id("Periods")
                batch_add_records("Periods", [[nid, pn, "Mo"]])
                st.rerun()

# --- TEACHER ---
def teacher_dashboard(period_id):
    user_email = st.session_state['user']['email']
    st.header(f"🍎 Giáo Viên: {st.session_state['user']['name']}")
    change_password_ui(user_email)
    
    classes = load_data("Classes")
    if classes.empty: return
    my_class = classes[classes['EmailGVCN'] == user_email]
    if my_class.empty:
        st.warning("Bạn chưa được phân công lớp.")
        return
    class_name = my_class.iloc[0]['TenLop']
    st.info(f"Lớp: {class_name}")

    # Load All Data Once
    users = load_data("Users")
    all_okrs = load_data("OKRs")
    all_reviews = load_data("FinalReviews")
    students = users[users['TenLop'] == class_name] if not users.empty and 'TenLop' in users.columns else pd.DataFrame()

    # --- METRICS & CHART ---
    if not students.empty:
        total_hs = len(students)
        submitted = 0
        ranks = {'Tốt': 0, 'Khá': 0, 'Đạt': 0, 'Chưa đạt': 0}
        
        for _, hs in students.iterrows():
            hs_okrs = all_okrs[(all_okrs['Email_HocSinh'] == hs['Email']) & (all_okrs['ID_Dot'] == period_id)] if not all_okrs.empty else pd.DataFrame()
            if not hs_okrs.empty:
                submitted += 1
                total_pct = sum([calculate_percent(r['ActualValue'], r['TargetValue']) for _, r in hs_okrs.iterrows()])
                avg = total_pct / len(hs_okrs) if len(hs_okrs) > 0 else 0
                r_txt, _ = get_rank(avg)
                ranks[r_txt] += 1
        
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Tổng HS", total_hs)
        m2.metric("Đã nhập", f"{submitted}/{total_hs}")
        m3.metric("Tỉ lệ", f"{round(submitted/total_hs*100, 1) if total_hs else 0}%")
        with m4:
            if submitted > 0:
                fig, ax = plt.subplots(figsize=(2, 2))
                ax.pie(list(ranks.values()), labels=list(ranks.keys()), autopct='%1.0f%%', textprops={'fontsize': 6})
                st.pyplot(fig)
    st.divider()

    # --- REPORT EXPORT ---
    period_name = "Học kỳ"
    p_df = load_data("Periods")
    if not p_df.empty:
        p_row = p_df[p_df['ID'] == period_id]
        if not p_row.empty: period_name = p_row.iloc[0]['TenDot']
    
    docx_class = create_class_report_docx(class_name, students, all_okrs, all_reviews, period_name, period_id)
    st.download_button("📥 XUẤT BÁO CÁO CẢ LỚP (.docx)", data=docx_class, file_name=f"BaoCaoLop_{class_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.divider()

    tab1, tab2 = st.tabs(["Danh sách & Duyệt", "Quản lý HS (Thêm/Import)"])

    with tab1: # LIST
        if students.empty:
            st.write("Chưa có học sinh.")
        else:
            c = st.columns([0.5, 2, 1.5, 1.5, 1.5])
            c[0].markdown("**STT**")
            c[1].markdown("**Họ Tên**")
            c[2].markdown("**Trạng thái L1**")
            c[3].markdown("**Trạng thái L2**")
            c[4].markdown("**Thao tác**")
            
            for idx, (_, hs) in enumerate(students.iterrows(), 1):
                hs_email = hs['Email']
                hs_okrs = all_okrs[(all_okrs['Email_HocSinh'] == hs_email) & (all_okrs['ID_Dot'] == period_id)] if not all_okrs.empty else pd.DataFrame()
                
                # Status
                if hs_okrs.empty: s1, b1 = "Chưa tạo", "badge-red"
                else:
                    s1, b1 = ("Chờ duyệt", "badge-yellow") if not hs_okrs[hs_okrs['TrangThai'] == 'ChoDuyet'].empty else ("Đã duyệt", "badge-green")
                
                has_rev = False
                if not all_reviews.empty:
                    rev = all_reviews[(all_reviews['Email_HocSinh'] == hs_email) & (all_reviews['ID_Dot'] == period_id)]
                    if not rev.empty and rev.iloc[0]['NhanXet_GV']: has_rev = True
                s2, b2 = ("Đã xong", "badge-green") if has_rev else ("Chưa xong", "badge-grey")
                
                has_del = False
                if not hs_okrs.empty and not hs_okrs[hs_okrs['DeleteRequest'] == 1].empty: has_del = True

                with st.container():
                    cols = st.columns([0.5, 2, 1.5, 1.5, 1.5])
                    cols[0].write(f"{idx}")
                    cols[1].write(hs['HoTen'] + (" ⚠️" if has_del else ""))
                    cols[2].markdown(f'<span class="status-badge {b1}">{s1}</span>', unsafe_allow_html=True)
                    cols[3].markdown(f'<span class="status-badge {b2}">{s2}</span>', unsafe_allow_html=True)
                    
                    c_act = cols[4]
                    if c_act.button("Chi tiết", key=f"v_{hs_email}"):
                        st.session_state['selected_hs'] = hs.to_dict()
                        st.rerun()
                    if c_act.button("🗑️ Xóa", key=f"quick_del_{hs_email}"):
                        delete_student_fully(hs_email)
                        st.rerun()

            st.divider()
            
            # --- DETAIL VIEW ---
            if 'selected_hs' in st.session_state:
                curr = st.session_state['selected_hs']
                st.markdown(f"### 📝 Chi tiết: {curr['HoTen']}")
                
                # Edit Info
                with st.expander("🛠️ Sửa thông tin (Email/Tên)"):
                    with st.form("edit_hs"):
                        new_e = st.text_input("Email", value=curr['Email'])
                        new_n = st.text_input("Tên", value=curr['HoTen'])
                        if st.form_submit_button("Lưu & Cập nhật"):
                            if new_e != curr['Email']:
                                if new_e in users['Email'].values: st.error("Email trùng!")
                                else:
                                    if update_student_email_cascade(curr['Email'], new_e):
                                        update_cell_value("Users", "Email", new_e, "HoTen", new_n)
                                        st.success("Xong!")
                                        st.session_state['selected_hs']['Email'] = new_e
                                        time.sleep(1)
                                        st.rerun()
                            elif new_n != curr['HoTen']:
                                update_cell_value("Users", "Email", curr['Email'], "HoTen", new_n)
                                st.rerun()

                # Report Download
                hs_okrs = all_okrs[(all_okrs['Email_HocSinh'] == curr['Email']) & (all_okrs['ID_Dot'] == period_id)] if not all_okrs.empty else pd.DataFrame()
                rev_gv, rev_ph = "", ""
                if not all_reviews.empty:
                    r_row = all_reviews[(all_reviews['Email_HocSinh'] == curr['Email']) & (all_reviews['ID_Dot'] == period_id)]
                    if not r_row.empty:
                        rev_gv = r_row.iloc[0]['NhanXet_GV']
                        rev_ph = r_row.iloc[0]['NhanXet_PH']

                docx_single = create_single_docx(curr['HoTen'], class_name, period_name, hs_okrs, rev_gv, rev_ph)
                st.download_button("📥 Tải phiếu kết quả (Word)", data=docx_single, file_name=f"KQ_{curr['HoTen']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                # OKR Items
                if hs_okrs.empty:
                    st.warning("Chưa có OKR.")
                else:
                    for i, row in hs_okrs.iterrows():
                        with st.container(border=True):
                            c1, c2, c3 = st.columns([4, 2, 2])
                            if row['DeleteRequest'] == 1: st.error("⚠️ Yêu cầu xóa")
                            c1.markdown(f"**O:** {row['MucTieu']}")
                            c1.text(f"KR: {row['KetQuaThenChot']}")
                            c2.metric("Target/Actual", f"{row['TargetValue']} / {row['ActualValue']} {row['Unit']}")
                            pct = calculate_percent(row['ActualValue'], row['TargetValue'])
                            c2.progress(min(pct/100, 1.0))
                            with c3:
                                if row['TrangThai'] == 'ChoDuyet':
                                    if st.button("✅ Duyệt", key=f"a_{row['ID']}"):
                                        update_cell_value("OKRs", "ID", row['ID'], "TrangThai", "DaDuyetMucTieu")
                                        st.rerun()
                                if row['DeleteRequest'] == 1:
                                    if st.button("🗑️ Chấp thuận xóa", key=f"d_{row['ID']}"):
                                        delete_record("OKRs", "ID", row['ID'])
                                        st.rerun()

                    st.write("---")
                    if rev_ph: st.info(f"🏠 Phụ huynh: {rev_ph}")
                    with st.form("final_rev"):
                        txt = st.text_area("Nhận xét GV", value=rev_gv)
                        if st.form_submit_button("Lưu đánh giá"):
                            upsert_final_review(curr['Email'], period_id, "NhanXet_GV", txt)
                            st.success("Đã lưu!")

    with tab2: # Manage HS
        # 1. Manual Add
        st.write("###### Thêm thủ công")
        with st.form("manual_add"):
            c1, c2 = st.columns(2)
            n_name = c1.text_input("Họ tên")
            n_email = c2.text_input("Email")
            if st.form_submit_button("Thêm HS"):
                if n_name and n_email:
                    existing = set(users['Email'].tolist()) if not users.empty else set()
                    if n_email in existing:
                        st.error("Email đã tồn tại!")
                    else:
                        batch_add_records("Users", [[n_email, "123", n_name, "HocSinh", class_name]])
                        st.success("Đã thêm!")
                        st.rerun()

        st.divider()
        # 2. Batch Import
        with st.expander("📥 Import Excel (Batch Upload - Chống lỗi 429)"):
            upl = st.file_uploader("Chọn file .xlsx", type=['xlsx'])
            if upl:
                try:
                    df_up = pd.read_excel(upl)
                    existing = set(users['Email'].tolist()) if not users.empty else set()
                    new_users, new_rels = [], []
                    count = 0
                    for _, r in df_up.iterrows():
                        e = str(r['Email']).strip()
                        n = str(r['HoTen']).strip()
                        if e and e not in existing:
                            new_users.append([e, "123", n, "HocSinh", class_name])
                            existing.add(e)
                            count += 1
                            if 'EmailPH' in r and pd.notna(r['EmailPH']):
                                ph = str(r['EmailPH']).strip()
                                new_rels.append([e, ph])
                    
                    if new_users: batch_add_records("Users", new_users)
                    if new_rels: batch_add_records("Relationships", new_rels)
                    st.success(f"Đã thêm {count} HS!")
                except Exception as e: st.error(f"Lỗi: {e}")

# --- STUDENT ---
def student_dashboard(period_id):
    user = st.session_state['user']
    st.header(f"🎒 {user['name']}")
    change_password_ui(user['email'])
    
    with st.expander("➕ Thêm OKR"):
        with st.form("add_okr"):
            mt = st.text_input("Mục tiêu")
            kr = st.text_input("Kết quả then chốt")
            c1, c2 = st.columns(2)
            tar = c1.number_input("Target", 0.1)
            unit = c2.text_input("Đơn vị", "Điểm")
            if st.form_submit_button("Lưu"):
                nid = get_next_id("OKRs")
                batch_add_records("OKRs", [[nid, user['email'], period_id, mt, kr, 0, 'ChoDuyet', '', '', '', tar, 0, unit, 0]])
                st.rerun()

    st.divider()
    all_okrs = load_data("OKRs")
    my_okrs = all_okrs[(all_okrs['Email_HocSinh'] == user['email']) & (all_okrs['ID_Dot'] == period_id)] if not all_okrs.empty else pd.DataFrame()
    
    if my_okrs.empty:
        st.info("Chưa có dữ liệu.")
    else:
        for _, row in my_okrs.iterrows():
            with st.container(border=True):
                c1, c2, c3 = st.columns([4, 2, 2])
                c1.markdown(f"**{row['MucTieu']}** - {row['KetQuaThenChot']}")
                status = row['TrangThai']
                if row['DeleteRequest'] == 1: status += " (Chờ xóa)"
                c1.caption(status)
                c2.metric("Tiến độ", f"{row['ActualValue']}/{row['TargetValue']}")
                
                with c3:
                    with st.popover("Báo cáo"):
                        with st.form(f"u_{row['ID']}"):
                            val = st.number_input("Đạt:", value=float(row['ActualValue']))
                            if st.form_submit_button("Lưu"):
                                update_cell_value("OKRs", "ID", row['ID'], "ActualValue", val)
                                st.rerun()
                    if row['TrangThai'] == 'ChoDuyet':
                        if st.button("🗑️ Xóa", key=f"d_{row['ID']}"):
                            delete_record("OKRs", "ID", row['ID'])
                            st.rerun()
                    elif row['DeleteRequest'] == 0:
                        if st.button("❌ Xin xóa", key=f"r_{row['ID']}"):
                            update_cell_value("OKRs", "ID", row['ID'], "DeleteRequest", 1)
                            st.rerun()

    st.divider()
    # Download My Report
    # Get class name
    u_df = load_data("Users")
    my_class = u_df[u_df['Email'] == user['email']].iloc[0]['TenLop'] if not u_df.empty else ""
    
    # Reviews
    all_reviews = load_data("FinalReviews")
    rev_gv, rev_ph = "", ""
    if not all_reviews.empty:
        rev = all_reviews[(all_reviews['Email_HocSinh'] == user['email']) & (all_reviews['ID_Dot'] == period_id)]
        if not rev.empty:
            rev_gv = rev.iloc[0]['NhanXet_GV']
            rev_ph = rev.iloc[0]['NhanXet_PH']
    
    # Get period Name
    p_df = load_data("Periods")
    p_name = ""
    if not p_df.empty:
        pr = p_df[p_df['ID'] == period_id]
        if not pr.empty: p_name = pr.iloc[0]['TenDot']

    docx = create_single_docx(user['name'], my_class, p_name, my_okrs, rev_gv, rev_ph)
    st.download_button("📥 Tải kết quả về máy", data=docx, file_name=f"KQ_{user['name']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    c1, c2 = st.columns(2)
    c1.info(f"Giáo viên: {rev_gv}")
    c2.success(f"Gia đình: {rev_ph}")

# --- PARENT ---
def parent_dashboard(period_id):
    user = st.session_state['user']
    st.header("👨‍👩‍👧‍👦 Phụ Huynh")
    
    rels = load_data("Relationships")
    if rels.empty: return
    my_child = rels[rels['Email_PhuHuynh'] == user['email']]
    if my_child.empty:
        st.warning("Chưa liên kết con.")
        return
        
    child_email = my_child.iloc[0]['Email_HocSinh']
    all_okrs = load_data("OKRs")
    df_okr = all_okrs[(all_okrs['Email_HocSinh'] == child_email) & (all_okrs['ID_Dot'] == period_id)] if not all_okrs.empty else pd.DataFrame()
    
    if not df_okr.empty:
        st.dataframe(df_okr[['MucTieu', 'KetQuaThenChot', 'TargetValue', 'ActualValue', 'TrangThai']])
        
        all_reviews = load_data("FinalReviews")
        curr_cmt = ""
        if not all_reviews.empty:
            r = all_reviews[(all_reviews['Email_HocSinh'] == child_email) & (all_reviews['ID_Dot'] == period_id)]
            if not r.empty: curr_cmt = r.iloc[0]['NhanXet_PH']
            
        with st.form("ph_cmt"):
            txt = st.text_area("Ý kiến gia đình", value=curr_cmt)
            if st.form_submit_button("Gửi"):
                upsert_final_review(child_email, period_id, "NhanXet_PH", txt)
                upsert_final_review(child_email, period_id, "DaGui_PH", 1)
                st.success("Đã gửi!")

# ==============================================================================
# 6. MAIN LOOP
# ==============================================================================
def main():
    if 'user' not in st.session_state:
        login_page()
    else:
        role = st.session_state['user']['role']
        with st.sidebar:
            st.write(f"Xin chào, **{st.session_state['user']['name']}**")
            if st.button("Đăng xuất"):
                del st.session_state['user']
                st.rerun()
            st.divider()
            
            periods = load_data("Periods")
            p_id = None
            if not periods.empty:
                if role != 'Admin': periods = periods[periods['TrangThai'] == 'Mo']
                if not periods.empty:
                    p_map = dict(zip(periods['TenDot'], periods['ID']))
                    p_name = st.selectbox("Chọn Đợt", list(p_map.keys()))
                    p_id = p_map[p_name]
                else: st.warning("Chưa có đợt hoạt động.")

        if p_id:
            if role == 'Admin': admin_dashboard(p_id)
            elif role == 'GiaoVien': teacher_dashboard(p_id)
            elif role == 'HocSinh': student_dashboard(p_id)
            elif role == 'PhuHuynh': parent_dashboard(p_id)
        elif role == 'Admin':
            admin_dashboard(0)

if __name__ == "__main__":
    main()
